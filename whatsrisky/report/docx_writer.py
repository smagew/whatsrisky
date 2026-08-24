"""DOCX report writer: priority-ordered findings with severity marking."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ..models import SEVERITY_ORDER, Finding, ScanReport, Severity
from ..util import clean_text

ID_PREFIX = {
    Severity.CRITICAL: "CRIT",
    Severity.HIGH: "HIGH",
    Severity.MEDIUM: "MED",
    Severity.LOW: "LOW",
    Severity.INFO: "INFO",
}

SEVERITY_MEANING = {
    Severity.CRITICAL: "Fix now. Exploitable with severe impact (RCE, auth bypass, live credential, data breach).",
    Severity.HIGH: "Fix before the next release. Exploitable with serious impact or a known high-severity CVE.",
    Severity.MEDIUM: "Plan remediation. Real weakness, but needs preconditions or has limited impact.",
    Severity.LOW: "Hardening / defense-in-depth. Schedule opportunistically.",
    Severity.INFO: "Informational. No direct security impact; hygiene and observations.",
}

MONO = "Menlo"


# --- low level docx helpers -------------------------------------------
def _shade(element, hex_fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    element.append(shd)


def _shade_cell(cell, hex_fill: str) -> None:
    _shade(cell._tc.get_or_add_tcPr(), hex_fill)


def _shade_paragraph(paragraph, hex_fill: str) -> None:
    _shade(paragraph._p.get_or_add_pPr(), hex_fill)


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def _field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def _code_block(doc, text: str, max_lines: int = 22) -> None:
    lines = clean_text(text).rstrip().splitlines()[:max_lines]
    if not lines:
        return
    for idx, line in enumerate(lines):
        para = doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.space_before = Pt(1 if idx else 4)
        fmt.space_after = Pt(4 if idx == len(lines) - 1 else 1)
        fmt.left_indent = Inches(0.15)
        fmt.keep_together = True
        _shade_paragraph(para, "F4F4F5")
        run = para.add_run(line[:220])
        run.font.name = MONO
        run.font.size = Pt(7.5)
        run.font.color.rgb = _rgb("212529")


def _label_value(doc, label: str, value: str, bold_value: bool = False) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(9.5)
    run2 = para.add_run(clean_text(value) or "-")
    run2.bold = bold_value
    run2.font.size = Pt(9.5)


def _badge(paragraph, severity: Severity) -> None:
    run = paragraph.add_run(f"  {severity.value}  ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = _rgb("FFFFFF")
    rpr = run._r.get_or_add_rPr()
    _shade(rpr, severity.color)


def _table(doc, headers: list[str], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for idx, title in enumerate(headers):
        cell = header_cells[idx]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(clean_text(title))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = _rgb("FFFFFF")
        _shade_cell(cell, "343A40")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def _fill_row(table, values: list[str], size: float = 9.0, widths: list[float] | None = None):
    row = table.add_row()
    for idx, value in enumerate(values):
        cell = row.cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(clean_text(value))
        run.font.size = Pt(size)
        if widths and idx < len(widths):
            cell.width = Inches(widths[idx])
    return row


# --- report sections --------------------------------------------------
def _document() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("whatsrisky  ·  page ").font.size = Pt(8)
    _field(footer, "PAGE", "1")
    return doc


def _cover(doc, report: ScanReport) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(90)
    run = title.add_run("Security Assessment Report")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = _rgb("1B1F23")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(clean_text(report.project_name))
    run.font.size = Pt(16)
    run.font.color.rgb = _rgb("495057")

    counts = report.counts()
    worst = next((s for s in SEVERITY_ORDER if counts[s]), None)
    verdict = doc.add_paragraph()
    verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    verdict.paragraph_format.space_before = Pt(18)
    if worst is not None:
        _badge(verdict, worst)
    run = verdict.add_run(f"  {report.verdict()}")
    run.bold = True
    run.font.size = Pt(12)

    score = doc.add_paragraph()
    score.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = score.add_run(f"Risk score {report.risk_score()}/100  ·  {len(report.findings)} findings")
    run.font.size = Pt(11)
    run.font.color.rgb = _rgb("6C757D")

    doc.add_paragraph().paragraph_format.space_before = Pt(24)
    meta = _table(doc, ["", ""], [1.8, 4.6])
    rows = [
        ("Project path", report.project_path),
        ("Git branch / commit", " / ".join(x for x in (report.git_branch, report.git_commit) if x) or "not a git repo"),
        ("Scan started", report.started_at),
        ("Scan finished", report.finished_at),
        ("Duration", f"{report.duration_s:.1f}s"),
        ("Scanners", ", ".join(f"{t.name} ({t.status})" for t in report.tools) or "-"),
    ]
    for label, value in rows:
        row = _fill_row(meta, [label, value], widths=[1.8, 4.6])
        row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_page_break()


def _toc(doc) -> None:
    doc.add_heading("Contents", level=1)
    note = doc.add_paragraph()
    run = note.add_run("Word shows the entries after you refresh fields (select all, then F9).")
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = _rgb("868E96")
    _field(doc.add_paragraph(), 'TOC \\o "1-2" \\h \\z \\u', "Right-click here and choose Update Field.")
    doc.add_page_break()


def _executive_summary(doc, report: ScanReport) -> None:
    doc.add_heading("1. Executive summary", level=1)
    counts = report.counts()
    total = len(report.findings)

    para = doc.add_paragraph()
    para.add_run("Verdict: ").bold = True
    run = para.add_run(report.verdict())
    run.bold = True
    worst = next((s for s in SEVERITY_ORDER if counts[s]), None)
    if worst is not None:
        run.font.color.rgb = _rgb(worst.color)
    doc.add_paragraph(
        f"{total} finding(s) were reported across {sum(1 for t in report.tools if t.ok)} scanner(s) "
        f"for `{report.project_name}`. The risk score is {report.risk_score()}/100, a weighted "
        "aggregate that saturates - it ranks projects, it does not measure them."
    )

    doc.add_heading("Findings by priority", level=2)
    table = _table(doc, ["Severity", "Count", "What it means / SLA"], [1.1, 0.7, 4.6])
    for severity in SEVERITY_ORDER:
        row = _fill_row(table, ["", str(counts[severity]), SEVERITY_MEANING[severity]], widths=[1.1, 0.7, 4.6])
        cell = row.cells[0]
        cell.text = ""
        run = cell.paragraphs[0].add_run(severity.value)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = _rgb("FFFFFF")
        _shade_cell(cell, severity.color)
        run_count = row.cells[1].paragraphs[0].runs[0]
        run_count.bold = True
        if counts[severity]:
            run_count.font.color.rgb = _rgb(severity.color)

    by_tool = report.counts_by_tool()
    if by_tool:
        doc.add_heading("Findings by scanner", level=2)
        headers = ["Scanner"] + [s.value for s in SEVERITY_ORDER] + ["Total"]
        table = _table(doc, headers)
        for tool, tool_counts in sorted(by_tool.items()):
            values = [tool] + [str(tool_counts[s]) for s in SEVERITY_ORDER] + [str(sum(tool_counts.values()))]
            _fill_row(table, values, size=8.5)

    top = [f for f in report.sorted_findings() if f.severity in (Severity.CRITICAL, Severity.HIGH)][:10]
    if top:
        doc.add_heading("Top priorities", level=2)
        for idx, finding in enumerate(top, 1):
            para = doc.add_paragraph(style="List Number" if idx == 1 else "List Number")
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(f"[{finding.severity.value}] ")
            run.bold = True
            run.font.color.rgb = _rgb(finding.severity.color)
            para.add_run(clean_text(f"{finding.title} — {finding.location} ({finding.tool})"))
    doc.add_page_break()


def _scope(doc, report: ScanReport) -> None:
    doc.add_heading("2. Scope and methodology", level=1)
    doc.add_paragraph(
        "Every scanner below ran against the project path with no changes to the working tree. "
        "Findings are normalized into one severity scale and de-duplicated per tool. "
        "Automated scanning is a floor, not a ceiling: it does not replace threat modelling, "
        "authenticated dynamic testing, or a manual review of business logic."
    )
    table = _table(doc, ["Scanner", "Covers", "Status", "Version", "Time", "Findings"], [1.0, 1.6, 0.8, 1.6, 0.6, 0.8])
    coverage = {
        "semgrep": "First-party source code (SAST)",
        "trivy": "Dependencies (CVE), IaC misconfig",
        "gitleaks": "Secrets in tree and git history",
        "claude": "LLM review of logic and auth",
    }
    for tool in report.tools:
        _fill_row(
            table,
            [
                tool.name,
                coverage.get(tool.name, "-"),
                tool.status,
                tool.version or "-",
                f"{tool.duration_s:.0f}s",
                str(len(tool.findings)),
            ],
            size=8.5,
            widths=[1.0, 1.6, 0.8, 1.6, 0.6, 0.8],
        )

    problems = [t for t in report.tools if t.status in ("missing", "error", "skipped")]
    if problems:
        doc.add_heading("Coverage gaps", level=2)
        doc.add_paragraph(
            "The following scanners did not contribute findings. Treat their area as UNSCANNED, "
            "not as clean."
        )
        for tool in problems:
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(f"{tool.name} — {tool.status}: ")
            run.bold = True
            para.add_run(clean_text(tool.message) or "no detail")

    commands = [t for t in report.tools if t.command]
    if commands:
        doc.add_heading("Commands executed", level=2)
        for tool in commands:
            _code_block(doc, f"$ {tool.command}")
    doc.add_page_break()


def _finding_block(doc, finding: Finding, ident: str) -> None:
    heading = doc.add_heading(level=3)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(clean_text(f"{ident}  {finding.title}"))
    run.font.color.rgb = _rgb(finding.severity.color)

    badge_para = doc.add_paragraph()
    badge_para.paragraph_format.space_after = Pt(6)
    badge_para.paragraph_format.keep_with_next = True
    _badge(badge_para, finding.severity)
    tail = f"  {finding.tool}"
    if finding.category:
        tail += f"  ·  {finding.category}"
    if finding.confidence:
        tail += f"  ·  confidence {finding.confidence}"
    run = badge_para.add_run(tail)
    run.font.size = Pt(9)
    run.font.color.rgb = _rgb("6C757D")

    _label_value(doc, "Location", finding.location, bold_value=True)
    if finding.rule_id:
        _label_value(doc, "Rule", finding.rule_id)
    if finding.package:
        version = finding.installed_version or "?"
        fixed = f" → fixed in {finding.fixed_version}" if finding.fixed_version else " → no fix available"
        _label_value(doc, "Package", f"{finding.package} {version}{fixed}")
    meta_bits = []
    if finding.cwe:
        meta_bits.append("CWE: " + ", ".join(finding.cwe[:6]))
    if finding.owasp:
        meta_bits.append("OWASP: " + ", ".join(finding.owasp[:3]))
    if finding.cvss:
        meta_bits.append(f"CVSS: {finding.cvss}")
    if meta_bits:
        _label_value(doc, "Classification", "  |  ".join(meta_bits))

    if finding.description:
        para = doc.add_paragraph()
        para.add_run("What is wrong").bold = True
        para.paragraph_format.space_after = Pt(2)
        for chunk in finding.description.split("\n\n"):
            if chunk.strip():
                body = doc.add_paragraph(clean_text(chunk).strip())
                body.paragraph_format.space_after = Pt(4)

    if finding.snippet:
        para = doc.add_paragraph()
        para.add_run("Evidence").bold = True
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.keep_with_next = True
        _code_block(doc, finding.snippet)

    if finding.remediation:
        para = doc.add_paragraph()
        para.add_run("How to fix").bold = True
        para.paragraph_format.space_after = Pt(2)
        doc.add_paragraph(clean_text(finding.remediation).strip())

    if finding.references:
        para = doc.add_paragraph()
        para.add_run("References").bold = True
        para.paragraph_format.space_after = Pt(2)
        for ref in finding.references[:5]:
            item = doc.add_paragraph(clean_text(ref), style="List Bullet")
            item.paragraph_format.space_after = Pt(1)
            for run in item.runs:
                run.font.size = Pt(8.5)
                run.font.color.rgb = _rgb("495057")

    ident_para = doc.add_paragraph()
    run = ident_para.add_run(f"id {finding.fingerprint}")
    run.font.size = Pt(7.5)
    run.font.color.rgb = _rgb("ADB5BD")


def _findings(doc, report: ScanReport, max_per_severity: int | None) -> None:
    doc.add_heading("3. Findings by priority", level=1)
    if not report.findings:
        doc.add_paragraph(
            "No findings were reported by the scanners that ran. Check section 2 for coverage gaps "
            "before reading this as a clean bill of health."
        )
        return

    grouped: dict[Severity, list[Finding]] = {s: [] for s in SEVERITY_ORDER}
    for finding in report.sorted_findings():
        grouped[finding.severity].append(finding)

    for severity in SEVERITY_ORDER:
        items = grouped[severity]
        if not items:
            continue
        heading = doc.add_heading(level=2)
        run = heading.add_run(f"{severity.value} ({len(items)})")
        run.font.color.rgb = _rgb(severity.color)
        note = doc.add_paragraph()
        run = note.add_run(SEVERITY_MEANING[severity])
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = _rgb("6C757D")

        shown = items if max_per_severity is None else items[:max_per_severity]
        for idx, finding in enumerate(shown, 1):
            _finding_block(doc, finding, f"{ID_PREFIX[severity]}-{idx:02d}")
        hidden = len(items) - len(shown)
        if hidden > 0:
            para = doc.add_paragraph()
            run = para.add_run(
                f"+ {hidden} more {severity.value} finding(s) omitted from this document; "
                "see the accompanying findings.json for the complete list."
            )
            run.italic = True
            run.font.size = Pt(9)


def _appendix(doc, report: ScanReport) -> None:
    ai_notes = [t for t in report.tools if t.name == "claude" and t.ok and t.message]
    diagnostics = [t for t in report.tools if t.stderr_tail or t.status in ("error", "missing")]
    if not ai_notes and not diagnostics:
        return
    doc.add_page_break()
    doc.add_heading("Appendix A. Scanner notes", level=1)
    for tool in ai_notes:
        doc.add_heading("AI reviewer summary", level=2)
        doc.add_paragraph(clean_text(tool.message))
    for tool in diagnostics:
        doc.add_heading(f"{tool.name} diagnostics", level=2)
        if tool.message:
            doc.add_paragraph(clean_text(tool.message))
        if tool.stderr_tail:
            _code_block(doc, tool.stderr_tail, max_lines=14)


def write_docx(report: ScanReport, out_path: Path, max_per_severity: int | None = None) -> Path:
    doc = _document()
    _cover(doc, report)
    _toc(doc)
    _executive_summary(doc, report)
    _scope(doc, report)
    _findings(doc, report, max_per_severity)
    _appendix(doc, report)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
